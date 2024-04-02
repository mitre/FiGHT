#!/usr/bin/env python
"""

NOTICE
This software was produced for the U. S. Government under Basic Contract No. W56KGU-18-D-0004, and
is subject to the Rights in Noncommercial Computer Software and Noncommercial Computer Software 
Documentation Clause 252.227-7014 (FEB 2012). Copyright (c) 2022 The MITRE Corporation.

This copyright notice must not be removed from this software, 
absent MITRE's express written permission.

"""
from os import listdir
from os.path import isfile, join
import sys
import traceback
from docx import Document
import logging
import csv
import re

headers = [
    "Procedure Examples", #0
    "Mitigations",        #1
    "Pre-Conditions",     #2
    "Critical Assets",    #3
    "Detection",          #4
    "Post-Conditions",     #5
    "References"          #6
]

fight_types = [
    "fight_technique",                           #0
    "fight_subtechnique",                        #1
    "fight_subtechnique_to_attack_technique",    #2
    "attack_technique_addendum",                 #3
    "attack_subtechnique_addendum",              #4
    "attack_technique_with_subs_with_addendums", #5
    "attack_technique_with_fight_subs",          #6
    "invalid_type"                               #6
]

docx_re_i = re.compile(re.escape('.docx'), re.IGNORECASE)
addendum_re_i = re.compile(re.escape("addendum"), re.IGNORECASE)

tid_re = re.compile("(^T\d\d\d\d).(\d\d\d)")
# FIGHT Technique TID
fight_tid_re = re.compile("^(FGT\d\d\d\d)(.*)$")
# FIGHT Subtechnique TID
fight_sub_tid_re = re.compile("^(FGT\d\d\d\d\.\d\d\d)(.*)$")
fight_sub_tid_re2 = re.compile("^(FGT\d\d\d\d)\.(\d\d\d)$")
# FIGHT Subtechnique added to existing ATT&CK
attack_fight_sub_tid_re = re.compile("(T\d\d\d\d\.5\d\d)(.*)")
attack_fight_sub_tid_re2 = re.compile("(T\d\d\d\d)\.(5\d\d)")
# ATT&CK Technique TID
attack_tid_re = re.compile("^(T\d\d\d\d)(.*)$")
attack_stid_re = re.compile("^(T\d\d\d\d.\d\d\d)")
# ATT&CK Subtechnique TID
attack_sub_tid_re = re.compile("^(T\d\d\d\d\.0\d\d)(.*)")
attack_sub_tid_re2 = re.compile("^(T\d\d\d\d)\.(0\d\d)")

types = [
    "fight_technique",                            #0
    "fight_subtechnique",                         #1
    "fight_subtechnique_to_attack_technique",     #2
    "attack_technique_addendum",                  #3
    "attack_subtechnique_addendum",               #4
    "attack_technique_with_subs_with_addendums",  #5
    "attack_technique_with_fight_subs",           #6
    "invalid_type"                                #6
]

def link_refs(text, offset=0):
    return text
    p = r'\[(\d+)\]'
    #r = r'<a href="#\1">\[\1\]</a>'
    r = r'[\[\1\]](#\1)'
    newtext = re.sub(p, r, text)
    return newtext

class FightDraft(object):
    """FiGHT Draft Object

    An object that represents the data parsed from a processed Word document.
    """

    def __init__(self):
        self.tid = {}
        self.name = ""
        self.description = ""
        self.subtechniques = []
        self.tactics = []
        self.architecture = []
        self.platforms = []
        self.access_type_req = []
        self.data_sources = []
        self.theoretical = False
        self.poc = False
        self.observed = False
        self.procedure_examples = []
        self.mitigations = {}
        self.pre_conditions = []
        self.critical_assets = []
        self.detections = {}
        self.post_conditions = []
        self.references = []

    def get_tid(self):
        return self.tid

    def proc_filename(self, filename):
        clean_filename = filename.lstrip().rstrip()
        clean_filename = addendum_re_i.sub("", clean_filename)
        fight_tid_m = fight_tid_re.search(clean_filename)
        fight_sub_tid_m = fight_sub_tid_re.search(clean_filename)
        attack_fight_sub_tid_m = attack_fight_sub_tid_re.search(clean_filename)
        attack_sub_tid_m = attack_sub_tid_re.search(clean_filename)
        attack_tid_m = attack_tid_re.search(clean_filename)
        if attack_fight_sub_tid_m:
            tid_value = attack_fight_sub_tid_m.groups()[0]
            tname = attack_fight_sub_tid_m.groups()[1]
            type_val = 2
        elif attack_sub_tid_m:
            tid_value = attack_sub_tid_m.groups()[0]
            tname = attack_sub_tid_m.groups()[1]
            type_val = 4
        elif attack_tid_m:
            tid_value = attack_tid_m.groups()[0]
            tname = attack_tid_m.groups()[1]
            type_val = 3
        elif fight_sub_tid_m:
            tid_value = fight_sub_tid_m.groups()[0]
            tname = fight_sub_tid_m.groups()[1]
            type_val = 1
        elif fight_tid_m:
            tid_value = fight_tid_m.groups()[0]
            tname = fight_tid_m.groups()[1]
            type_val = 0
        else:
            logging.warning(f"DOCX PARSE FATAL ERROR! UNRECOGNIZED TID FORMAT for {filename}!'")
            self.set_type(6)
            return False
        logging.debug(f'Found type {type_val} for {tid_value}')
        self.set_type(type_val)
        self.set_tid(tid_value)
        self.set_name(tname)
        return True

    def set_tid(self, tid):
        self.tid = tid

    def get_type(self):
        return self.type

    def set_type(self, this_type):
        if this_type < len(fight_types):
            self.type = fight_types[this_type]
        else:
            raise Exception(f"Improper type value {this_type} given")

    def get_name(self):
        return self.name

    def set_name(self, name):
        if ":" in name:
            fight_name = name.split(":")[1]
        elif ";" in name:
            fight_name = name.split(";")[1]
        else:
            fight_name = name
        fight_name = docx_re_i.sub("", fight_name)
        self.name = fight_name.rstrip().lstrip()

    def get_description(self):
        return self.description

    def set_description(self, description):
        self.description = description

    def get_subtechniques(self):
        return self.subtechniques

    def add_subtechniques(self, subtechnique):
        if subtechnique == "None" or subtechnique == "N/A":
            return
        else:
            self.subtechniques.append(subtechnique)

    def get_tactics(self):
        return self.tactics

    def add_tactics(self, tactic):
        self.tactics.append(tactic)

    def get_architecture(self):
        return self.architecture

    def get_architeture_str(self):
        if self.architecture:
            return ",".join(self.architecture)
        else:
            return ""

    def add_architecture(self, architecture):
        self.architecture.append(architecture)

    def get_platforms(self):
        return self.platforms

    def add_platform(self, platform):
        self.platforms.append(platform)

    def get_access_type_req(self):
        return self.access_type_req

    def get_access_type_req_str(self):
        if self.access_type_req:
            return ",".join(self.access_type_req)
        else:
            return ""

    def add_access_type_req(self, access_type):
        self.access_type_req.append(access_type)

    def get_data_sources(self):
        return self.data_sources

    def add_data_source(self, data_source):
        self.data_sources.append(data_source)

    def check_theoretical(self):
        return self.theoretical

    def set_theoretical(self):
        self.theoretical = True

    def unset_theoretical(self):
        self.theoretical = False

    def get_class_md(self):
        if self.type == fight_types[5] or self.type == fight_types[6]:
            return "This is an observed behavior in Enterprise networks"
        elif self.type == fight_types[3] or self.type == fight_types[4]:
            return  "This is an observed behavior in Enterprise networks, and is theoretical in context of 5G systems."
        if self.check_theoretical():
            if self.type == fight_types[3] or self.type == fight_types[4]:
                return  "This is an observed behavior in Enterprise networks, and is theoretical in context of 5G systems."
            else:
                return  "This is a theoretical behavior in context of 5G systems."
        elif self.check_poc():
            return "This a 5G relevant behavior that has been demonstrated in a successful proof of concept"
        elif self.check_observed():
            return "Observed in earlier 3GPP generations and expected in 5G."
        # If not specified, assume theoretical
        else:
            return  "This is a theoretical behavior"

    def get_addendum_class_md(self):
        if self.check_theoretical():
            return  "This is a 5G specific, theoretical behavior variant\r\n"
        elif self.check_poc():
            return "This a 5G specific behavior variant that has been demonstrated in a proof of concept\r\n"
        elif self.check_observed():
            return "This a 5G specific behavior variant that has been observed in the wild\r\n"
        # If not specified, assume theoretical
        else:
            return  "##### This is a theoretical 5G specific behavior variant\r\n"

    def check_poc(self):
        return self.poc

    def set_poc(self):
        self.poc = True

    def unset_poc(self):
        self.poc = False

    def check_observed(self):
        return self.observed

    def set_observed(self):
        self.observed = True

    def unset_observed(self):
        self.observed = False

    def get_procedure_examples(self):
        return self.procedure_examples

    def get_procedure_examples_dict(self, offset=0):
        myarray = []
        for procedure_ex in self.procedure_examples:
            name = procedure_ex[0]
            desc = link_refs(procedure_ex[1], offset)
            myarray.append({"Name": name, "Description": desc})
        return myarray

    def get_procedure_examples_md(self):
        if self.check_theoretical():
            markdown = "#### Possible Implementation Examples\r\n"
        elif self.check_poc():
            markdown = "#### Possible Implementation Examples\r\n"
        elif self.observed:
            markdown = "#### Procedure Examples\r\n"
        else:
            logging.warning(f"CAUTION!  {self.tid} IS NOT SPECIFIED THEORETICAL / POC / OBSERVED!")
            # ASSUME THEORETICAL
            markdown = "#### Possible Implementation Examples\r\n"
        if len(self.procedure_examples):
            for procedure in self.procedure_examples:
                header = procedure[0]
                body = procedure[1]
                section = f"##### {header}\r\n{body}"
                markdown += f"{section}\r\n"
            return markdown
        else:
            return f"{markdown}\r\nNone."

    def add_procedure_examples(self, name, example):
        self.procedure_examples.append([name, example])

    def get_mitigations(self):
        return self.mitigations

    def add_mitigations(self, fgmid, use):
        logging.info(f"Gathering Mitigation row for {self.tid}: {fgmid}: {use}")
        if fgmid in self.mitigations:
            logging.warning(f"Duplicate FGMID {fgmid} found for FGTID {self.tid}, SKIPPING")
        else:
            self.mitigations[fgmid] = use

    def get_pre_conditions(self):
        return self.pre_conditions

    def get_pre_conditions_dict(self):
        myarray = []
        for pre_condition in self.pre_conditions:
            name = pre_condition[0]
            desc = pre_condition[1]
            myarray.append({"Name": name, "Description": desc})
        return myarray

    def get_pre_conditions_md(self):
        markdown = "#### Pre Conditions\r\n"
        if len(self.pre_conditions):
            for pre_condition in self.pre_conditions:
                header = pre_condition[0]
                body = pre_condition[1]
                section = f"##### {header}\r\n{body}"
                markdown += f"{section}\r\n"
            return markdown
        else:
            return f"{markdown}\r\nNone."

    def add_pre_conditions(self, name, pre_condition):
        self.pre_conditions.append([name, pre_condition])

    def get_critical_assets_dict(self):
        myarray = []
        for asset in self.critical_assets:
            name = asset[0]
            desc = asset[1]
            myarray.append({"Name": name, "Description": desc})
        return myarray

    def get_critical_assets(self):
        return self.critical_assets

    def add_critical_assets(self, name, critical_asset):
        self.critical_assets.append([name, critical_asset])

    def get_detections(self):
        return self.detections

    def add_detections(self, fgdsid, detects):
        logging.info(f"Gathering Data Source row for {self.tid}: {fgdsid}: {detects}")
        if fgdsid in self.detections:
            logging.warning(f"Duplicate FGDSID {fgdsid} found for FGTID {self.tid}, SKIPPING")
        else:
            self.detections[fgdsid] = detects

    def get_post_conditions(self):
        return self.post_conditions

    def get_post_conditions_dict(self):
        myarray = []
        for post_condition in self.post_conditions:
            name = post_condition[0]
            desc = post_condition[1]
            myarray.append({"Name": name, "Description": desc})
        return myarray

    def get_post_conditions_md(self):
        markdown = "#### Post Conditions\r\n"
        if len(self.post_conditions):
            for post_condition in self.post_conditions:
                header = post_condition[0]
                body = post_condition[1]
                section = f"##### {header}\r\n{body}"
                markdown += f"{section}\r\n"
            return markdown
        else:
            return f"{markdown}\r\nNone."

    def add_post_conditions(self, name, post_condition):
        self.post_conditions.append([name, post_condition])

    def get_references(self):
        return self.references

    def get_references_dict(self):
        myarray = []
        count = 1
        for reference in self.references:
            text = reference[0]
            url = reference[1]
            link = f"<a name=\"{count}\"> \[{count}\] </a> [{text}]({url})"
            myarray.append(link)
            count += 1
        return myarray

    def get_references_md(self):
        markdown = "#### References\r\n"
        number = 1
        if len(self.references):
            for reference in self.references:
                text = reference[0]
                url = reference[1]
                link = f"[{text}]({url})"
                markdown += f"##### {number} {link}\r\n"
                number += 1
            return markdown
        else:
            return f"{markdown}\r\nNone."

    def add_references(self, text, url):
        logging.debug(f"Adding reference {text} | {url} for {self.tid}:{self.name}")
        self.references.append([text, url])

    def get_text(self):
        text = f"{self.tid}: {self.name}"
        text += f"{self.description}"
        header_i = 0
        text += ", ".join(self.subtechniques)

        for header in headers:
            text += "-" * 40 + "\n"
            text += f"{header}\n"
            text += "-" * 40 + "\n"
            if header_i == 0 or header_i == "":
                for row in self.get_procedure_examples():
                    text += f"{row[0]}: {row[1]}\n"
            elif header_i == 1:
                mitigations = self.get_mitigations()
                for fgmid in mitigations:
                    text += f"{fgmid}: {mitigations[fgmid]}\n"
            elif header_i == 2:
                for row in self.get_pre_conditions():
                    text += f"{row[0]}: {row[1]}\n"
            elif header_i == 3:
                for row in self.get_critical_assets():
                    text += f"{row[0]}: {row[1]}\n"
            elif header_i == 4:
                detections = self.get_detections()
                for fgdsid in detections:
                    text += f"{fgdsid}: {detections[fgdsid]}\n"
            elif header_i == 5:
                for row in self.get_post_conditions():
                    text += f"{row[0]}: {row[1]}\n"
            elif header_i == 6:
                for row in self.get_post_conditions():
                    text += f"{row[0]}: {row[1]}\n"
            header_i += 1
        return text

class WordImporter(object):
    """Word Importer - import Microsoft Word FiGHT Template Document and convert to YAML format.
    Needs to be able to import a folder filled with MS Word FiGHT documents and convert to YAML
    for deployment to the website.
    """
    def __init__(self, directory, filename):
        def get_parts(func, text):
            parts = text.split(":")
            if len(parts) > 1:
                for part in parts[1].split(","):
                    logging.info(f"Got {part} for {tid}")
                    func(part)
        self.full_filename = join(directory, filename)
        logging.info(f"Processing {filename}")
        self.draft = None
        if isfile(self.full_filename) and ".docx" in self.full_filename:
            self.draft = FightDraft()
            document = Document(self.full_filename)
            result = self.draft.proc_filename(filename)
            if not result:
                return
            tid = self.draft.get_tid()
            tname = self.draft.get_name()
            num_tables = len(document.tables)
            logging.info(f"{tid} {tname} has num tables: {num_tables}")
            offset = 0
            if document.tables[0].rows[0].cells[0].text == "Date" \
                    and document.tables[0].rows[0].cells[1].text == "Who":
                logging.info(f"Processing {tid}: Found version control table, skipping")
                offset = 1
            line_no = 0
            description = ""
            indesc = False
            for paragraph in document.paragraphs:
                # Skip the first line
                logging.debug(f"{tid}: {paragraph.text}")
                #self.debug_dump_draft()
                if line_no == 0:
                    pass
                if "Description:" in paragraph.text:
                    description += paragraph.text.replace("Description:","")
                    description += "\r\n"
                    indesc = True
                elif "Labelling" in paragraph.text:
                    logging.info(f"Found labelling for {tid}")
                    indesc = False
                    self.draft.set_description(description)
                elif indesc:
                    description += f"{paragraph.text}\r\n"
                elif "Sub-Techniques" in paragraph.text:
                    get_parts(self.draft.add_subtechniques, paragraph.text)
                elif "Applicable Tactics" in paragraph.text:
                    get_parts(self.draft.add_tactics, paragraph.text)
                elif "Architecture Segment" in paragraph.text:
                    get_parts(self.draft.add_architecture, paragraph.text)
                elif "Platforms" in paragraph.text:
                    get_parts(self.draft.add_platform, paragraph.text)
                elif "Access type required" in paragraph.text:
                    get_parts(self.draft.add_access_type_req, paragraph.text)
                elif "Data Sources" in paragraph.text:
                    get_parts(self.draft.add_data_source, paragraph.text)
                elif re.search("theoretical", paragraph.text, re.IGNORECASE) \
                        and re.search("Proof of Concept", paragraph.text, re.IGNORECASE) \
                            and re.search("Observed", paragraph.text, re.IGNORECASE):
                    value = paragraph.text.split(":")[1]
                    if re.search("theoretical", value, re.IGNORECASE):
                        self.draft.set_theoretical()
                    elif re.search("proof of concept", value, re.IGNORECASE):
                        self.draft.set_poc()
                    elif re.search("observed", value, re.IGNORECASE):
                        self.draft.set_observed()
                    else:
                        logging.info(f"Unknown answer for T/PoC/O: {value} in document {self.full_filename}")
                elif "doNotParse" in paragraph.text:
                    break
                else:
                    header_c = 0
                    for header in headers:
                        if header in paragraph.text and header_c+offset < num_tables:
                            logging.debug(f"Processing {tid}: Found {header} for table #{header_c}")
                            rownum = len(document.tables[header_c+offset].rows)
                            logging.debug(f"Document {tid} has {rownum} rows in table #{header_c}")
                            self.grab_table(paragraph.text, document.tables[header_c+offset].rows)
                        header_c += 1
                line_no += 1
            logging.info(f"Finished processing {tid}:{tname}")

    def debug_dump_draft(self):
        doc_text = self.draft.get_text()
        tid = self.draft.get_tid()
        logging.debug(f"TEXT DUMP FOR {tid}")
        logging.debug(f"{doc_text}")
        logging.debug(f"END TEXT DUMP FOR {tid}")

    def grab_table(self, para_text, doc_rows):
        row_i = 0
        for row in doc_rows:
            logging.debug(f"\tProcessing row {row_i}")
            if row_i == 0:
                row_i += 1
                continue
            this_name = row.cells[0].text.rstrip()
            this_desc = row.cells[1].text.rstrip()
            logging.debug(f"\t\t{this_name} | {this_desc}")
            if this_name == "Name" and this_desc == "Description":
                row_i += 1
                continue
            elif this_name == "If known":
                row_i += 1
                continue
            elif this_name == "Specific example if known":
                row_i += 1
                continue
            elif this_name == "" and this_desc == "":
                row_i += 1
                continue
            elif "name that will appear" in this_name:
                row_i += 1
                continue
            logging.debug(f"\t{this_name}: {this_desc}")
            if headers[0] in para_text:
                self.draft.add_procedure_examples(this_name, this_desc)
            elif headers[1] in para_text:
                self.draft.add_mitigations(this_name, this_desc)
            elif headers[2] in para_text:
                self.draft.add_pre_conditions(this_name, this_desc)
            elif headers[3] in para_text:
                self.draft.add_critical_assets(this_name, this_desc)
            elif headers[4] in para_text:
                self.draft.add_detections(this_name, this_desc)
            elif headers[5] in para_text:
                self.draft.add_post_conditions(this_name, this_desc)
            elif headers[6] in para_text:
                urls = ""
                for para in row.cells[1].paragraphs:
                    urls += GetParagraphText(para)
                logging.debug(f"FOUND URLS:{urls}")
                if urls:
                    self.draft.add_references(this_name, urls)
                else:
                    self.draft.add_references(this_name, this_desc)
            row_i += 1

    def get_draft(self):
        return self.draft

def GetParagraphText(paragraph):
    def GetTag(element):
        return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))
    text = ''
    runCount = 0
    for child in paragraph._p:
        tag = GetTag(child)
        if tag == "w:r":
            text += paragraph.runs[runCount].text
            runCount += 1
        if tag == "w:hyperlink":
            for subChild in child:
                if GetTag(subChild) == "w:r":
                    text += subChild.text
    return text

def get_drafts(directory):
    drafts = {}
    for filename in listdir(directory):
        try:
            file_importer = WordImporter(directory, filename)
            draft = file_importer.get_draft()
            if draft:
                if draft.get_type() == fight_types[5]:
                    logging.info(f"Invalid type found for {filename}, skipping")
                    continue
                draft_tid = draft.get_tid()
                logging.info(f"Got {draft_tid}")
                if draft_tid in drafts:
                    drafts[draft_tid].append(draft)
                else:
                    drafts[draft_tid] = [draft]
            else:
                logging.warn(f"Did not get back a draft for {filename}")
        except Exception:
            logging.info(traceback.format_exc())
            logging.info(sys.exc_info()[2])
    return drafts

def dump_tables(drafts):
    # First, go file by file
    objects = {}
    for header in headers:
        objects[header] = []
    csv_file = open(f"static/fight_matrix/change_tracking_files/draft_status.csv", "w")
    csv_writer = csv.writer(csv_file)
    header_row = ["FGTID", "Name"]
    for header in headers:
        header_row.append(header)
    csv_writer.writerow(header_row)
    for tid in drafts:
        # THIS IGNORES SUBSEQUENT ADDENDUMS!
        draft = drafts[tid][0]
        name = draft.get_name()
        no_space_name = name.replace(" ", "_")
        outfile = open(f"static/fight_matrix/change_tracking_files/{tid}_{no_space_name}.txt", "w")
        outfile.write("="*80 + "\n")
        outfile.write(f"{tid}: {name}\n")
        outfile.write(f"Architecture Segment: {drafts[tid][0].get_architeture_str()}\r\n")
        this_row = [tid, no_space_name]
        header_i = 0
        for header in headers:
            outfile.write("-"*40 + "\n")
            outfile.write(f"{header}\n")
            outfile.write("-"*40 + "\n")
            if header_i == 0:
                for row in draft.get_procedure_examples():
                    outfile.write(f"{row[0]}: {row[1]}\n")
                    objects[header].append([tid, name, row[0], row[1]])
                if len(draft.get_procedure_examples()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            elif header_i == 1:
                mitigations = draft.get_mitigations()
                for fgmid in mitigations:
                    outfile.write(f"{fgmid}: {mitigations[fgmid]}\n")
                    objects[header].append([tid, name, fgmid, mitigations[fgmid]])
                if len(draft.get_mitigations()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            elif header_i == 2:
                for row in draft.get_pre_conditions():
                    outfile.write(f"{row[0]}: {row[1]}\n")
                    objects[header].append([tid, name, row[0], row[1]])
                if len(draft.get_pre_conditions()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            elif header_i == 3:
                for row in draft.get_critical_assets():
                    outfile.write(f"{row[0]}: {row[1]}\n")
                    objects[header].append([tid, name, row[0], row[1]])
                if len(draft.get_critical_assets()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            elif header_i == 4:
                detections = draft.get_detections()
                for fgdsid in detections:
                    outfile.write(f"{fgdsid}: {detections[fgdsid]}\n")
                    objects[header].append([tid, name, fgdsid, detections[fgdsid]])
                if len(draft.get_detections()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            elif header_i == 5:
                for row in draft.get_post_conditions():
                    outfile.write(f"{row[0]}: {row[1]}\n")
                    objects[header].append([tid, name, row[0], row[1]])
                if len(draft.get_post_conditions()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            elif header_i == 6:
                for row in draft.get_references():
                    outfile.write(f"{row[0]}: {row[1]}\n")
                    objects[header].append([tid, name, row[0], row[1]])
                if len(draft.get_references()) > 0:
                    this_row.append("Y")
                else:
                    this_row.append("N")
            header_i += 1
        csv_writer.writerow(this_row)
    csv_file.close()
    for header in objects:
        csv_file = open(f"{header}.csv", "w")
        csv_writer = csv.writer(csv_file)
        for row in objects[header]:
            csv_writer.writerow(row)
        csv_file.close()

if __name__ == "__main__":
    logging.basicConfig(filename='word_importer.log', level=logging.DEBUG)
    directory = sys.argv[1]
    drafts = get_drafts(directory)
    dump_tables(drafts)
